"""
Document generator for creating professional DOCX files.
"""
from pathlib import Path
from typing import List, Optional
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from app.config import OUTPUT_DIR
from app.utils import setup_logger, sanitize_filename

logger = setup_logger(__name__)


class DocumentGenerator:
    """Generates professional DOCX documents."""
    
    def __init__(self):
        """Initialize document generator."""
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self) -> None:
        """Configure document styles."""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
    
    def add_title(self, title: str) -> None:
        """Add document title."""
        heading = self.doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.size = Pt(24)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    
    def add_subtitle(self, subtitle: str) -> None:
        """Add document subtitle."""
        para = self.doc.add_paragraph(subtitle)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.size = Pt(12)
            run.font.italic = True
            run.font.color.rgb = RGBColor(64, 64, 64)
    
    def add_metadata(self, goal: str, assumptions: List[str]) -> None:
        """Add metadata section with goal and assumptions."""
        self.add_heading("Document Information", level=1)
        self.doc.add_paragraph("Goal:", style='List Bullet')
        goal_para = self.doc.add_paragraph(goal)
        goal_para.paragraph_format.left_indent = Inches(0.5)
        if assumptions:
            self.doc.add_heading("Key Assumptions", level=2)
            for i, assumption in enumerate(assumptions, 1):
                self.doc.add_paragraph(assumption, style='List Bullet')
    
    def add_heading(self, text: str, level: int = 1) -> None:
        """Add section heading."""
        heading = self.doc.add_heading(text, level=level)
        for run in heading.runs:
            if level == 1:
                run.font.color.rgb = RGBColor(0, 51, 102)
            else:
                run.font.color.rgb = RGBColor(0, 102, 153)
    
    def add_paragraph(self, text: str, bold: bool = False, italic: bool = False) -> None:
        """Add paragraph."""
        for line in text.split("\n"):
            if not line.strip():
                self.doc.add_paragraph("")
                continue

            para = self.doc.add_paragraph(line)
            for run in para.runs:
                run.font.bold = bold
                run.font.italic = italic
    
    def add_bullet_list(self, items: List[str]) -> None:
        """Add bulleted list."""
        for item in items:
            self.doc.add_paragraph(item, style='List Bullet')
    
    def add_numbered_list(self, items: List[str]) -> None:
        """Add numbered list."""
        for item in items:
            self.doc.add_paragraph(item, style='List Number')
    
    def add_section_with_content(
        self,
        heading: str,
        content: str,
        subsections: Optional[dict] = None,
    ) -> None:
        """
        Add a complete section with heading and content.
        
        Args:
            heading: Section title
            content: Main content
            subsections: Optional dict of {subsection_title: subsection_content}
        """
        self.add_heading(heading, level=1)
        self.add_paragraph(content)
        
        if subsections:
            for sub_title, sub_content in subsections.items():
                self.add_heading(sub_title, level=2)
                self.add_paragraph(sub_content)
    
    def add_table(
        self,
        headers: List[str],
        rows: List[List[str]],
    ) -> None:
        """Add table to document."""
        table = self.doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
        for row_idx, row_data in enumerate(rows, 1):
            row_cells = table.rows[row_idx].cells
            for col_idx, cell_data in enumerate(row_data):
                row_cells[col_idx].text = str(cell_data)
    
    def add_page_break(self) -> None:
        """Add page break."""
        self.doc.add_page_break()
    
    def add_footer(self) -> None:
        """Add footer with timestamp."""
        section = self.doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for run in footer_para.runs:
            run.font.size = Pt(8)
            run.font.italic = True
    
    def save(self, filename: str) -> str:
        """
        Save document to file.
        
        Args:
            filename: Filename (without .docx)
            
        Returns:
            Full path to saved file
        """
        self.add_footer()
        safe_name = sanitize_filename(filename)
        file_path = OUTPUT_DIR / f"{safe_name}.docx"
        
        try:
            self.doc.save(str(file_path))
            logger.info(f"Document saved to {file_path}")
            return str(file_path)
        
        except Exception as e:
            logger.error(f"Failed to save document: {e}")
            raise


def create_document_from_content(
    title: str,
    subtitle: str,
    goal: str,
    assumptions: List[str],
    content: str,
    filename: str = "report",
) -> str:
    """
    Create a complete document from content.
    
    Args:
        title: Document title
        subtitle: Document subtitle
        goal: Document goal
        assumptions: List of assumptions
        content: Main content
        filename: Output filename (without .docx)
        
    Returns:
        Path to saved document
    """
    logger.info(f"Creating document: {title}")
    
    generator = DocumentGenerator()
    generator.add_title(title)
    generator.add_subtitle(subtitle)
    generator.add_metadata(goal, assumptions)
    generator.add_page_break()
    generator.add_heading("Content", level=1)
    generator.add_paragraph(content)
    return generator.save(filename)


def get_generator() -> DocumentGenerator:
    """Factory function to get DocumentGenerator instance."""
    return DocumentGenerator()
